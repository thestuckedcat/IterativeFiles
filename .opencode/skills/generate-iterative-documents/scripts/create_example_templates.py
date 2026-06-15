#!/usr/bin/env python3
"""Create portable example templates for the iterative document skill."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
RED = RGBColor(255, 0, 0)


def red_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    return run


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def control_paragraph(doc, expression):
    red_run(doc.add_paragraph(), "{%p " + expression + " %}")


def dynamic_heading(doc, prefix, variable, level):
    p = doc.add_heading(level=level)
    p.add_run(prefix)
    red_run(p, "{{ " + variable + " }}")


def header(doc, placeholder):
    section = doc.sections[0]
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = red_run(p, placeholder)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    red_run(p, "{{ TITLE_CN }}").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    red_run(p, "{{ TITLE_EN }}").bold = True


def field(doc, label, variable):
    p = doc.add_paragraph()
    p.add_run(label)
    red_run(p, "{{ " + variable + " }}")


def make_srs(path):
    doc = Document()
    header(doc, "{{ HEADER }}")
    title_block(doc)
    heading(doc, "1 简介")
    field(doc, "1.1 目的\n", "PURPOSE")
    field(doc, "1.2 范围\n", "SCOPE")
    field(doc, "1.3 总体概述\n", "OVERVIEW")
    field(doc, "1.4 软件功能\n", "SOFTWARE_FUNCTIONS")
    field(doc, "1.5 设计约束\n", "DESIGN_CONSTRAINTS")
    heading(doc, "2 需求说明")
    control_paragraph(doc, "for req in requirements")
    dynamic_heading(doc, "功能需求：", "req.title", 2)
    field(doc, "AR号：", "req.ar")
    for label, key in [("介绍", "introduction"), ("输入", "input"), ("处理", "processing"),
                       ("输出", "output"), ("外部依赖", "external_dependencies"),
                       ("性能需求", "performance"), ("质量需求", "quality"),
                       ("可测试性", "testability")]:
        control_paragraph(doc, "if req." + key)
        heading(doc, label, 3)
        field(doc, label + "：", "req." + key)
        control_paragraph(doc, "endif")
    control_paragraph(doc, "endfor")
    heading(doc, "3 非功能需求")
    for label, key in [("3.1 性能需求", "PERFORMANCE_SUMMARY"), ("3.2 质量需求", "QUALITY_SUMMARY"),
                       ("3.3 可测试性", "TESTABILITY_SUMMARY")]:
        control_paragraph(doc, "if " + key)
        heading(doc, label, 2)
        field(doc, label + "\n", key)
        control_paragraph(doc, "endif")
    heading(doc, "4 修订记录")
    table = doc.add_table(rows=4, cols=4)
    for cell, text in zip(table.rows[0].cells, ["日期", "修订版本", "修改描述", "作者"]):
        cell.text = text
    red_run(table.rows[1].cells[0].paragraphs[0], "{%tr for rev in revisions %}")
    vals = ["rev.date", "rev.version", "rev.description", "rev.author"]
    for cell, val in zip(table.rows[2].cells, vals):
        red_run(cell.paragraphs[0], "{{ " + val + " }}")
    red_run(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")
    doc.save(path)


def make_sd(path):
    doc = Document()
    header(doc, "{{ HEADER }}")
    title_block(doc)
    field(doc, "迭代：", "ITERATION_REQUIREMENT")
    heading(doc, "1 简介")
    field(doc, "1.1 目的\n", "PURPOSE")
    field(doc, "1.2 范围\n", "SCOPE")
    field(doc, "1.3 软件功能\n", "SOFTWARE_FUNCTIONS")
    heading(doc, "3 需求详细设计")
    control_paragraph(doc, "for design in designs")
    dynamic_heading(doc, "", "design.title", 2)
    field(doc, "AR号：", "design.ar")
    for label, key in [("背景", "background"), ("详细设计", "detailed_design"),
                       ("场景约束", "constraints"), ("对外接口描述", "external_interface_link")]:
        control_paragraph(doc, "if design." + key)
        heading(doc, label, 3)
        red_run(doc.add_paragraph(), "{{ design." + key + " }}")
        control_paragraph(doc, "endif")
    control_paragraph(doc, "endfor")
    heading(doc, "4 修订记录")
    table = doc.add_table(rows=4, cols=6)
    labels = ["日期", "修订版本", "AR号", "修改章节", "修改描述", "作者"]
    for cell, text in zip(table.rows[0].cells, labels):
        cell.text = text
    red_run(table.rows[1].cells[0].paragraphs[0], "{%tr for rev in revisions %}")
    vals = ["rev.date", "rev.version", "rev.ar", "rev.section", "rev.description", "rev.author"]
    for cell, val in zip(table.rows[2].cells, vals):
        red_run(cell.paragraphs[0], "{{ " + val + " }}")
    red_run(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")
    doc.save(path)


def make_stc(path):
    wb = Workbook()
    stats = wb.active
    stats.title = "统计"
    cases = wb.create_sheet("测试用例")
    wb.create_sheet("缺陷记录&测试报告")
    headers = ["测试用例编号", "用例类型", "测试项", "测试用例标题", "重要级别", "预置条件",
               "输入", "操作步骤", "预期结果", "用例执行情况", "备注"]
    cases.append(headers)
    for cell in cases[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF7")
    stats.append(["统计项", "数量"])
    stats.append(["用例总数", 0])
    stats.append(["H", 0])
    stats.append(["M", 0])
    stats.append(["L", 0])
    wb.save(path)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_srs(ASSETS / "srs-template.docx")
    make_sd(ASSETS / "sd-template.docx")
    make_stc(ASSETS / "stc-template.xlsx")
    print(f"Created templates in {ASSETS}")


if __name__ == "__main__":
    main()
